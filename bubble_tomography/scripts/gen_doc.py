# -*- coding: utf-8 -*-
"""生成《气泡三维多相机层析重建系统》介绍与使用说明 Word 文档。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\songyuchen\Seafile\Code\BubbleandPIV\气泡三维多相机层析重建系统_介绍与使用说明.docx"

doc = Document()

# ---------- 颜色与基础字体 ----------
ACCENT = RGBColor(0x2C, 0x6E, 0xB5)   # 蓝色
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x55, 0x55, 0x55)
CODE_BG = "F2F3F5"
HEADER_FILL = "2C6EB5"
ALT_FILL = "EDF3FA"

CN_FONT = "微软雅黑"
EN_FONT = "Calibri"

def set_cn(run, font=CN_FONT, size=None, bold=None, color=None, italic=None):
    run.font.name = EN_FONT
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.font.italic = italic

# ---------- 页面设置 A4 ----------
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.4)
sec.right_margin = Cm(2.4)

# ---------- 默认正文样式 ----------
normal = doc.styles['Normal']
normal.font.name = EN_FONT
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), CN_FONT)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.25
pf.space_after = Pt(4)

# 标题样式
for hs, sz in [('Heading 1', 16), ('Heading 2', 13.5), ('Heading 3', 11.5)]:
    st = doc.styles[hs]
    st.font.name = EN_FONT
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = ACCENT
    st._element.rPr.rFonts.set(qn('w:eastAsia'), CN_FONT)
    st.paragraph_format.space_before = Pt(10)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True

def add_para(text="", size=10.5, bold=False, color=DARK, align=None,
             after=4, before=0, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_cn(r, size=size, bold=bold, color=color, italic=italic)
    return p

def add_runs(p, segments):
    """segments: list of (text, dict-of-kwargs)"""
    for text, kw in segments:
        r = p.add_run(text)
        set_cn(r, **kw)
    return p

def add_bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        set_cn(r, size=10.5, bold=True, color=ACCENT)
    r = p.add_run(text)
    set_cn(r, size=10.5)
    return p

def add_number(text, bold_lead=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        set_cn(r, size=10.5, bold=True, color=ACCENT)
    r = p.add_run(text)
    set_cn(r, size=10.5)
    return p

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=DARK, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_cn(r, size=size, bold=bold, color=color)

def add_table(headers, rows, widths=None, header_fill=HEADER_FILL,
              alt_fill=ALT_FILL, header_color=RGBColor(0xFF, 0xFF, 0xFF)):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=header_color, size=9.5,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr[i], header_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            set_cell_text(cells[ci], str(val), size=9.5)
            if alt_fill and ri % 2 == 1:
                shade_cell(cells[ci], alt_fill)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def add_code(code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), CODE_BG)
    pPr.append(shd)
    # 边框
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '6')
    left.set(qn('w:space'), '6'); left.set(qn('w:color'), 'C8CDD4')
    pbdr.append(left)
    pPr.append(pbdr)
    for line in code.strip('\n').split('\n'):
        r = p.add_run(line + '\n')
        set_cn(r, font="Consolas", size=9, color=RGBColor(0x1A, 0x1A, 0x1A))
    return p

def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'BFD0E0')
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(2)

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "右键此处，选择“更新域”以生成目录。"
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run._r.append(t); run._r.append(fld3)

# =====================================================================
# 封面
# =====================================================================
for _ in range(3):
    doc.add_paragraph()
add_para("气泡三维多相机层析重建系统", size=26, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
add_para("Bubble Tomographic Reconstruction System", size=14, bold=False,
         color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
add_para("介绍与使用说明", size=18, bold=True, color=DARK,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
hrule()
add_para("软件版本：v2.0（构建 20260627，包版本 1.2.0）", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
add_para("运行环境：Windows（便携版免安装）/ Python 3.13 源码运行", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
add_para("文档生成日期：2026-07-11", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
for _ in range(6):
    doc.add_paragraph()
add_para("© 多相流三维测量研究组  ·  许可协议：MIT License", size=9.5,
         color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# =====================================================================
# 目录
# =====================================================================
add_para("目  录", size=16, bold=True, color=ACCENT, after=8)
add_toc()
doc.add_page_break()

# =====================================================================
# 1. 软件概述
# =====================================================================
doc.add_heading('1. 软件概述', level=1)

doc.add_heading('1.1 软件名称与版本', level=2)
add_para("本软件全称“气泡三维多相机层析重建系统”（Bubble Tomographic "
         "Reconstruction System），是一套面向多相流实验测量的桌面级科学计算软件。")
add_para("当前版本：v2.0（构建号 20260627，Python 包版本 1.2.0）。")

doc.add_heading('1.2 适用领域', level=2)
add_para("软件面向气液两相流、沸腾传热、鼓泡塔、微通道等实验场景，主要解决以下"
         "科学问题：")
add_bullet("透明/半透明容器内上升气泡的三维形态、尺寸与空间分布重建；")
add_bullet("基于多相机同步拍摄的层析体素场重建与三维点云提取；")
add_bullet("示踪粒子三维重建与 Tomographic PIV（体互相关）三维速度场测量；")
add_bullet("连续时刻粒子轨迹追踪（PTV），获取拉格朗日速度场；")
add_bullet("单相机光线追踪几何重建（Silhouette + Snell 折射修正），适合双视角气泡测量。")

doc.add_heading('1.3 核心能力', level=2)
add_para("软件以“多相机标定 → 图像预处理 → 三维重建 → 速度场 → 可视化导出”"
         "为主线，集成了完整的实验测量工作流：")
add_table(
    ["能力", "说明"],
    [
        ["多相机标定", "支持棋盘格/圆点/非对称圆点/体标定板，3~N 相机联合标定"],
        ["气泡图像预处理", "背景去除、去畸变、二值化分割、柔化边缘投影计算"],
        ["层析重建", "MART / SMART / Conv-SMART 三种算法，光线追踪投影"],
        ["三维点云输出", "PLY / PCD / OBJ 格式，兼容 CloudCompare、MeshLab、ParaView"],
        ["单相机3D重建", "Silhouette 轮廓法 + Snell 折射光线追踪修正"],
        ["Tomographic PIV", "三维粒子重建 + 体互相关速度场"],
        ["二维 PIV", "单相机对/双视图互相关速度场测量"],
        ["PTV 粒子跟踪", "四帧前后、最近邻、松弛法、Shake-The-Box 多种跟踪"],
        ["批量时间序列", "多时刻图像批量加载、重建与速度场计算，时间点滑块切换"],
    ],
    widths=[4.0, 13.0],
)

doc.add_heading('1.4 技术架构', level=2)
add_para("软件采用 Python 3.13 + PyQt5 实现图形界面，计算核心基于 NumPy / "
         "SciPy / OpenCV / scikit-image / Matplotlib，并以 PyInstaller 打包为"
         "Windows 便携版。整体由以下模块构成：")
add_table(
    ["模块目录", "功能"],
    [
        ["main.py", "主入口：GUI / Demo / PIV-Demo / CLI 启动与调度"],
        ["calibration/", "多相机标定器（MultiCameraCalibrator）"],
        ["utils/", "气泡图像预处理、通用图像编辑工具"],
        ["mart/", "层析重建算法（MART / SMART / Conv-SMART）与光线追踪器"],
        ["raytrace/", "单相机光线追踪三维重建（Silhouette + Snell）"],
        ["particles/", "三维粒子重建、二维 PIV 互相关、速度场计算"],
        ["ptv/", "粒子跟踪测速（多种 PTV 算法 + 拉格朗日速度场）"],
        ["visualization/", "三维可视化、点云/体素/速度场导出"],
        ["gui/", "PyQt5 主窗口与全部交互界面"],
    ],
    widths=[3.6, 13.4],
)

# =====================================================================
# 2. 安装与运行
# =====================================================================
doc.add_heading('2. 安装与运行', level=1)

doc.add_heading('2.1 系统要求', level=2)
add_bullet("操作系统：Windows 10 / 11（便携版）；源码同样支持 Linux / macOS。", )
add_bullet("便携版对硬件无特殊要求，普通带摄像机的实验电脑即可运行。")
add_bullet("如需编译或从源码运行：Python 3.13，建议 8 GB 以上内存（层析重建内存随网格分辨率立方增长）。")

doc.add_heading('2.2 便携版运行（推荐，免安装）', level=2)
add_para("项目已打包为 PyInstaller 便携版，无需安装 Python 或任何依赖即可直接运行：")
add_bullet("启动文件：项目根目录下的 BubbleTomography.exe；")
add_bullet("运行库目录：同目录的 _internal（含 Python 运行时、PyQt5、OpenCV、NumPy、"
           "SciPy、Matplotlib 等），请勿单独删除或移动；")
add_bullet("使用方法：直接双击 BubbleTomography.exe 即可启动图形界面；")
add_bullet("迁移到新电脑：复制“BubbleTomography.exe + _internal”两个顶层项即可，"
           "或复制整个 dist\\BubbleTomography 发布文件夹。")
add_para("若双击无反应，可在该目录打开命令行执行 .\\BubbleTomography.exe 查看系统级错误提示；"
         "如被杀毒软件拦截，可将文件夹移至本地可信路径（如 C:\\Tools\\BubbleTomography）后运行。",
         color=GREY, size=9.5, italic=True)

doc.add_heading('2.3 源码方式运行', level=2)
add_para("安装依赖：", bold=False)
add_code("cd bubble_tomography\n"
         "pip install -r requirements.txt\n"
         "# 或手动安装：\n"
         "pip install numpy opencv-contrib-python scipy matplotlib scikit-image tqdm pyqt5")
add_para("启动图形界面：")
add_code("python main.py --gui")

doc.add_heading('2.4 启动参数与运行模式', level=2)
add_table(
    ["参数", "说明"],
    [
        ["（无参数）", "默认启动图形界面（GUI）"],
        ["--gui", "启动 PyQt5 图形界面（推荐日常使用）"],
        ["--demo", "运行气泡层析重建演示（使用合成数据，输出至 demo_output/）"],
        ["--piv-demo", "运行 Tomographic PIV 演示（输出至 piv_demo_output/）"],
        ["--cli", "命令行模式"],
        ["-v / --verbose", "输出详细日志"],
    ],
    widths=[3.5, 13.5],
)

# =====================================================================
# 3. 界面总览
# =====================================================================
doc.add_heading('3. 界面总览（GUI 模块导航）', level=1)
add_para("图形界面左侧为分组导航栏（带图标与悬浮效果的“iOS 风格”按钮），右侧为"
         "对应模块的 QStackedWidget 内容区。导航按功能分为四大组：")

doc.add_heading('3.1 导航结构', level=2)
add_table(
    ["分组", "模块", "主要用途"],
    [
        ["通用图像模块", "图像处理", "单/多图浏览、裁剪、灰度化、滤波、标尺测量等通用编辑"],
        ["三维图像模块", "相机标定", "多相机标定参数求解与精度评估"],
        ["三维图像模块", "气泡重建", "多视角气泡图像层析重建、点云提取、时间点切换"],
        ["三维图像模块", "单相机3D重建", "基于光线追踪的双视角气泡三维几何重建"],
        ["PIV 模块", "三维 PIV", "示踪粒子三维重建 + 体互相关速度场"],
        ["PIV 模块", "二维 PIV", "双视图二维互相关速度场测量"],
        ["PIV 模块", "PTV", "多帧粒子跟踪测速（拉格朗日轨迹）"],
        ["AI 辅助", "AI 辅助模型 / 本地模型", "调用 AI 模型辅助分析（按部署情况可用）"],
    ],
    widths=[3.2, 4.3, 9.5],
)

doc.add_heading('3.2 通用交互说明', level=2)
add_bullet("科学图像查看器支持鼠标左键拖动旋转、滚轮缩放、右键标尺测量像素距离；")
add_bullet("左侧文件树可批量加载实验图像，按文件夹自动归类（子文件夹=时间点）；")
add_bullet("重计算任务在后台线程执行，界面不卡顿，进度可在状态栏查看；")
add_bullet("结果图可一键导出为 1080p 高清图片。")

# =====================================================================
# 4. 详细使用流程
# =====================================================================
doc.add_heading('4. 详细使用流程', level=1)

doc.add_heading('4.1 相机标定', level=2)
add_para("标定是后续三维重建精度的基础。打开“相机标定”模块：")
add_number("添加相机：为每个相机指定 ID（如 cam1、cam2…）；")
add_number("加载标定图像：每个相机至少 3 张、推荐 5–10 张不同角度/距离的标定板图像；")
add_number("设置标定板参数：选择类型、内角点数量、方格边长（mm）；")
add_number("执行标定：软件自动检测角点/圆心并联合求解内外参，输出重投影误差报告。")
add_para("支持的标定板类型：")
add_table(
    ["类型", "参数值", "说明"],
    [
        ["棋盘格", "checkerboard", "最常用，黑白交替方格"],
        ["对称圆点阵", "circles", "规则排列的圆形"],
        ["非对称圆点阵", "acircles", "交错排列的圆形，精度更高"],
        ["体标定板点阵", "volume_dots", "亮圆点体标定板，支持少量缺失/编码点"],
    ],
    widths=[3.5, 3.5, 10.0],
)
add_para("标定图像拍摄建议：标定板应覆盖图像不同区域（中心、边缘），保持平整无反光，"
         "图像清晰对焦。GUI 同时支持自动检测与手动指定标定板原点。")

doc.add_heading('4.2 气泡图像预处理', level=2)
add_para("在“气泡重建”模块中，先对原始图像做预处理，得到适合层析投影的二值/"
         "灰度投影数据：")
add_bullet("背景去除：使用无气泡参考图做差，抑制背景噪声；")
add_bullet("去畸变：依据标定得到的内参/畸变系数校正图像；")
add_bullet("二值化分割：Otsu 等自适应阈值分割气泡区域；")
add_bullet("形态学去噪：开闭运算去除孤立噪声点；")
add_bullet("投影计算：生成柔化边缘（soft_edge）投影，供光线追踪使用。")
add_para("命令行等价示例（参考）：", size=10, after=2)
add_code("from utils import BubbleImageProcessor\n"
         "processor = BubbleImageProcessor(\n"
         "    background_method='reference',\n"
         "    threshold_method='otsu',\n"
         "    morph_operations=True)\n"
         "projections = processor.prepare_projection_data(\n"
         "    bubble_images, camera_intrinsics,\n"
         "    reference_images=ref, projection_type='soft_edge')")

doc.add_heading('4.3 三维层析重建（MART / SMART / Conv-SMART）', level=2)
add_para("软件提供三种统一的层析重建算法，均基于光线追踪的乘法代数重建技术（ART）：")
add_table(
    ["算法", "特点", "适用场景"],
    [
        ["MART", "逐光线乘法更新，经典 MART，实现简单", "常规气泡重建，参数直观"],
        ["SMART", "同步乘法更新，体素→光线反向索引，数值更稳定", "相机多、视角覆盖广、收敛更稳"],
        ["Conv-SMART", "卷积核（PSF）+ FFT 加速，内存降低 10–100 倍", "大网格 / 高分辨率重建"],
    ],
    widths=[3.0, 7.0, 7.0],
)
add_para("在“气泡重建”模块中选择算法并设置参数后执行重建，软件将：")
add_bullet("基于光线追踪，从各相机像素发射光线并计算与体素交点权重；")
add_bullet("迭代更新体素场，达到最大迭代次数或收敛阈值后停止；")
add_bullet("基于 Marching Cubes 提取气泡表面点云（按体素阈值）。")
add_para("关键重建参数（ReconstructionConfig）：", size=10, after=2)
add_code("config = ReconstructionConfig(\n"
         "    grid_size=(64, 64, 64),      # 重建网格分辨率\n"
         "    domain_size=(20, 20, 20),    # 重建域物理尺寸 (mm)\n"
         "    relaxation_factor=0.5,       # 松弛因子 mu (0~1)\n"
         "    max_iterations=50,           # 最大迭代次数\n"
         "    voxel_threshold=0.1,         # 表面提取阈值\n"
         "    ray_sample_step=0.2,         # 光线采样步长 (mm)\n"
         "    algorithm='MART')           # 'MART' / 'SMART' / 'ConvSMART'")
add_para("Conv-SMART 专属参数：conv_kernel_size（卷积核尺寸）、psf_type（gaussian/"
         "tophat/empirical）、psf_sigma、use_fft_convolution。")

doc.add_heading('4.4 单相机三维重建（光线追踪）', level=2)
add_para("“单相机3D重建”模块面向透明容器内气泡的双视角几何重建，算法移植自 "
         "MATLAB raytrace_main.m，流程为：二值化 → Silhouette 轮廓三维重建 → "
         "面法向初始化 → 光线追踪迭代修正 → 曲面构建。采用 Silhouette 方法结合 "
         "Snell 定律折射修正，适合仅双相机视角的气泡三维几何测量。")

doc.add_heading('4.5 结果可视化与导出', level=2)
add_para("重建完成后可在“结果可视化”模块查看并导出：")
add_bullet("三维点云（可旋转查看，支持配色按高度/法向）；")
add_bullet("体素切片（沿 x/y/z 轴多切片对比）；")
add_bullet("投影对比（原始投影 vs 重建投影）；")
add_bullet("综合报告图、批量结果概览。")
add_para("导出格式：PLY / PCD / OBJ 点云、NPY 体素数据，可直接导入 CloudCompare、"
         "MeshLab、ParaView、Blender 等软件进一步处理。")

doc.add_heading('4.6 三维 PIV（Tomographic PIV）', level=2)
add_para("在“三维 PIV”模块中完成示踪粒子三维重建与速度场测量：")
add_number("加载多相机示踪粒子图像，按相机 ID 匹配；")
add_number("粒子检测与多视角三角测量 + 外极线约束匹配，得到三维粒子坐标；")
add_number("对连续两帧粒子做体互相关，计算三维速度场；")
add_number("查看速度矢量图、速度切面、信噪比（SNR）统计，导出 VTK 速度场。")
add_para("体互相关关键参数（CorrelationConfig）：interrogation_size（ interrogation "
         "窗口尺寸 mm）、overlap_ratio（重叠率）、subpixel_refinement（亚像素细化）、"
         "peak_threshold、max_displacement、median_filter。")

doc.add_heading('4.7 二维 PIV', level=2)
add_para("“二维 PIV”模块面向单相机对/双视图的平面互相关速度场测量。支持互相关窗口大小"
         "设置、矢量样式（比例、颜色模式、箭头形态）调整，并支持鼠标框选/多边形选择"
         "“无粒子区域”作为排除掩膜，提升矢量质量。")

doc.add_heading('4.8 PTV 粒子跟踪测速', level=2)
add_para("“PTV”模块实现连续帧的三维粒子轨迹追踪与拉格朗日速度场计算，内置多种经典算法：")
add_table(
    ["算法", "说明"],
    [
        ["四帧前后跟踪", "Forward-Backward Tracking，利用前后四帧提高轨迹连续性"],
        ["最近邻跟踪", "Nearest Neighbor，按位移阈值匹配相邻帧粒子"],
        ["松弛法跟踪", "Relaxation Method，迭代优化全局匹配"],
        ["Shake-The-Box", "STB（Schanz et al., 2016），高密度 3D-PTV 迭代优化方法"],
    ],
    widths=[4.0, 13.0],
)
add_para("软件基于多帧粒子位置序列构建轨迹（Track），计算位移与平均速度，输出拉格朗日"
         "速度场与轨迹统计。")

doc.add_heading('4.9 批量时间序列处理', level=2)
add_para("GUI 支持从根目录批量加载多时刻图像，自动按子文件夹识别时间点：")
add_code("root_dir/\n"
         "├── t000/\n"
         "│   ├── cam1.png\n"
         "│   ├── cam2.png\n"
         "│   └── cam3.png\n"
         "├── t001/ ...\n"
         "└── t002/ ...")
add_bullet("每个子文件夹对应一个时间点，文件夹名作为时间标识；")
add_bullet("子文件夹内图像名需包含对应相机 ID（如 cam1.png）以自动匹配；")
add_bullet("右侧面板底部“时间点滑块”可快速切换查看各时刻的重建/速度场结果；")
add_bullet("支持批量重建与批量 PIV，结果可整体导出。")

# =====================================================================
# 5. 参数调优
# =====================================================================
doc.add_heading('5. 重建参数说明与调优', level=1)
add_para("下表汇总 MART / SMART / Conv-SMART 核心参数及调优建议：")
add_table(
    ["参数", "默认值", "说明", "调优建议"],
    [
        ["grid_size", "(64,64,64)", "三维网格分辨率", "增大提高精度，但计算量立方增长"],
        ["domain_size", "(20,20,20)", "重建域物理尺寸 (mm)", "应略大于气泡群实际范围"],
        ["relaxation_factor", "0.5", "松弛因子 μ (0~1)", "0.1~0.3 更稳但慢；0.5~0.8 快但易发散"],
        ["max_iterations", "50", "最大迭代次数", "通常 20~50 次即可收敛"],
        ["voxel_threshold", "0.1", "表面提取阈值", "依重建值分布调整"],
        ["ray_sample_step", "0.2", "光线采样步长 (mm)", "越小越精确但越慢"],
        ["conv_kernel_size", "5", "卷积核尺寸（Conv-SMART）", "影响 PSF 近似范围"],
        ["psf_sigma", "1.0", "高斯 PSF 的 sigma", "依光学模糊程度调整"],
    ],
    widths=[3.3, 2.6, 4.1, 7.0],
)

# =====================================================================
# 6. 输出格式
# =====================================================================
doc.add_heading('6. 输出文件格式', level=1)
add_table(
    ["格式", "扩展名", "兼容软件"],
    [
        ["PLY", ".ply", "MeshLab, CloudCompare, ParaView"],
        ["PCD", ".pcd", "CloudCompare, PCL Viewer"],
        ["OBJ", ".obj", "Blender, MeshLab"],
        ["NPY", ".npy", "Python (numpy.load)"],
        ["VTK", ".vtk", "ParaView, 三维速度场可视化"],
        ["PNG", ".png", "通用图像，1080p 高清结果图"],
    ],
    widths=[2.5, 2.5, 12.0],
)

# =====================================================================
# 7. FAQ
# =====================================================================
doc.add_heading('7. 常见问题（FAQ）', level=1)
add_para("Q1：标定重投影误差太大怎么办？", bold=True, color=ACCENT, after=2)
add_para("确保标定图像清晰、标定板平整、角点完整检测。尝试增加标定图像数量（10+ 张），"
         "覆盖图像不同区域。", before=0)
add_para("Q2：重建结果出现伪影？", bold=True, color=ACCENT, after=2)
add_para("降低松弛因子（0.1~0.3），增加迭代次数，减小光线采样步长；确保相机角度覆盖"
         "足够（相邻相机间隔建议 ≤45°）。", before=0)
add_para("Q3：重建速度太慢？", bold=True, color=ACCENT, after=2)
add_para("减小网格分辨率（如 32³）、增大光线采样步长、减少迭代次数。MART 计算复杂度约为 "
         "O(N_rays × N_voxels × N_iterations)；大数据量优先使用 Conv-SMART。", before=0)
add_para("Q4：便携版在新电脑上无法启动？", bold=True, color=ACCENT, after=2)
add_para("确认 BubbleTomography.exe 与 _internal 在同一目录；若被杀软拦截，移至本地可信"
         "路径运行；双击无反应时从命令行启动查看错误。", before=0)
add_para("Q5：批量处理时图像未被识别？", bold=True, color=ACCENT, after=2)
add_para("检查子文件夹内图像文件名是否包含相机 ID（如 cam1.png），且层级结构符合"
         "“根目录/时间点/相机图像”规范。", before=0)

# =====================================================================
# 8. 技术参考
# =====================================================================
doc.add_heading('8. 技术参考', level=1)
add_para("若本软件对您的研究有帮助，可在论文中引用相关方法：")
add_bullet("MART / ART 算法：Gordon, R., Bender, R., & Herman, G. T. (1970). "
           "Algebraic reconstruction techniques (ART) for three-dimensional electron "
           "microscopy and X-ray photography. Journal of Theoretical Biology, 29(3), 471-481.")
add_bullet("Shake-The-Box：Schanz, D., Gesemann, S., & Schröder, A. (2016). "
           "Shake-The-Box: Lagrangian particle tracking at high particle image densities. "
           "Experiments in Fluids, 57(5), 70.")
add_bullet("单相机光线追踪：基于 Silhouette 轮廓法 + Snell 定律折射修正（移植自 "
           "raytrace_main.m）。")

# =====================================================================
# 9. 许可
# =====================================================================
doc.add_heading('9. 许可与致谢', level=1)
add_para("本软件以 MIT License 开源发布，可自由用于学术研究与教学。使用过程中如产生改进"
         "或衍生工作，欢迎反馈以便持续改进。")
add_para("致谢：感谢多相流三维测量研究组在算法移植、界面设计与实验验证方面的贡献。",
         color=GREY, size=9.5, italic=True)

# =====================================================================
# 页脚（页码）
# =====================================================================
def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第 ")
    set_cn(r, size=9, color=GREY)
    # PAGE field
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    r2 = p.add_run(); r2._r.append(fld1); r2._r.append(instr); r2._r.append(fld2)
    set_cn(r2, size=9, color=GREY)
    r3 = p.add_run(" 页")
    set_cn(r3, size=9, color=GREY)

add_page_number_footer(doc.sections[0])

doc.save(OUT)
print("Saved:", OUT)
